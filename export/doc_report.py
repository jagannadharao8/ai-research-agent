from docx import Document
from docx.shared import Pt, Inches
import io

def generate_markdown(query, answer, sources):
    md = f"# Research Report: {query}\n\n"
    md += "## AI Answer\n"
    md += f"{answer}\n\n"
    
    if sources:
        md += "## Sources Referenced\n"
        for doc in sources:
            citation = doc.get("citation", "")
            title = doc.get("title", "Untitled")
            url = doc.get("url", "")
            if url:
                md += f"- [{citation}] [{title}]({url})\n"
            else:
                md += f"- [{citation}] {title}\n"
                
    return md

def generate_docx(query, answer, sources):
    doc = Document()
    
    # Title
    doc.add_heading(f'Research Report', 0)
    doc.add_heading(query, 1)
    
    # Answer
    doc.add_heading('AI Answer', level=2)
    doc.add_paragraph(answer)
    
    # Sources
    if sources:
        doc.add_heading('Sources Referenced', level=2)
        for d in sources:
            citation = d.get("citation", "")
            title = d.get("title", "Untitled")
            url = d.get("url", "")
            text = f"[{citation}] {title}"
            if url:
                text += f" ({url})"
            doc.add_paragraph(text, style='List Bullet')
            
    # Save to memory stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
